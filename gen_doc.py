#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate 第五组_任务分工分配表.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Global style defaults ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.25
style.paragraph_format.space_after = Pt(4)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    hs.font.bold = True

# ── Helper: set table header shading ──
def set_header_shading(cell, color_hex='1F4E79'):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, font_size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def set_col_widths(table, widths):
    """Set column widths (list of Cm)."""
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

# ════════════════════════════════════════════════
#  TITLE
# ════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('项目任务分工分配表')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run('多人脸采集素材自动化清洗')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ════════════════════════════════════════════════
#  第一部分：项目基本信息
# ════════════════════════════════════════════════
doc.add_heading('一、项目基本信息', level=1)

table1 = doc.add_table(rows=10, cols=2)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(table1, [Cm(5), Cm(11)])

rows_data = [
    ('组号', '第五组'),
    ('组长', '彭仁希（23050910）'),
    ('组员1', '王启炀（23050911）'),
    ('组员2', '余野（23050912）'),
    ('组员3', '于枫硕（23050913）'),
    ('组员4', '侯贺升（23050730）'),
    ('组员5', '吴迪（23050732）'),
    ('项目经理', '王启炀'),
    ('需求分析师', '于枫硕'),
    ('开发人员', '彭仁希'),
]

for i, (label, value) in enumerate(rows_data):
    row = table1.rows[i]
    set_cell_text(row.cells[0], label, bold=True, font_size=Pt(10.5),
                  align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[1], value, font_size=Pt(10.5),
                  align=WD_ALIGN_PARAGRAPH.LEFT)
    if i % 2 == 0:
        set_header_shading(row.cells[0], 'D6E4F0')
        set_header_shading(row.cells[1], 'D6E4F0')
    else:
        set_header_shading(row.cells[0], 'EBF3FB')
        set_header_shading(row.cells[1], 'EBF3FB')

# Add test person as extra row
row = table1.add_row()
set_cell_text(row.cells[0], '测试人员', bold=True, font_size=Pt(10.5),
              align=WD_ALIGN_PARAGRAPH.LEFT)
set_cell_text(row.cells[1], '余野', font_size=Pt(10.5),
              align=WD_ALIGN_PARAGRAPH.LEFT)
set_header_shading(row.cells[0], 'D6E4F0')
set_header_shading(row.cells[1], 'D6E4F0')

doc.add_paragraph()

# ════════════════════════════════════════════════
#  第二部分：成员角色与主要任务
# ════════════════════════════════════════════════
doc.add_heading('二、成员角色与主要任务', level=1)

members = [
    ('彭仁希', '23050910', '组长 / 开发人员', [
        '技术方案设计与架构搭建',
        '人脸检测核心模块开发（人脸检测 + 对齐 + 裁剪）',
        '项目代码审查与合并',
        '协助编写技术文档',
    ]),
    ('王启炀', '23050911', '项目经理', [
        '制定项目计划与里程碑',
        '组织例会与进度跟踪',
        '协调各模块接口对接',
        '风险管控与问题升级',
        '答辩PPT与演示准备',
    ]),
    ('余野', '23050912', '测试人员', [
        '编写测试计划与测试用例',
        '功能测试与回归测试',
        '性能测试（大量图片处理效率）',
        'Bug报告与跟踪闭环',
        '最终质量评估报告',
    ]),
    ('于枫硕', '23050913', '需求分析师', [
        '需求调研与竞品分析',
        '编写需求规格说明书（SRS）',
        '定义系统功能模块与数据流图',
        '制定功能验收标准',
        '编写用户操作手册',
    ]),
    ('侯贺升', '23050730', '组员', [
        '多源人脸数据采集模块开发（爬虫/下载器）',
        '图像去重算法实现（感知哈希/特征匹配）',
        '数据去重与初步清洗脚本',
    ]),
    ('吴迪', '23050732', '组员', [
        '图像质量评估模块（模糊度检测、光照检测、遮挡检测）',
        '清洗规则引擎开发',
        '系统前端界面/可视化面板开发',
        '系统集成与打包',
    ]),
]

for name, sid, role, tasks in members:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{name}（{sid}）— {role}')
    run.font.bold = True
    run.font.size = Pt(11.5)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    for task in tasks:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(task)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ════════════════════════════════════════════════
#  第三部分：功能模块与负责人
# ════════════════════════════════════════════════
doc.add_heading('三、功能模块与负责人', level=1)

table2 = doc.add_table(rows=9, cols=4)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(table2, [Cm(3.5), Cm(5.5), Cm(2.5), Cm(4)])

headers = ['功能模块', '模块说明', '负责人', '协作人']
for j, h in enumerate(headers):
    set_cell_text(table2.rows[0].cells[j], h, bold=True, font_size=Pt(10),
                  color=RGBColor(0xFF, 0xFF, 0xFF))
    set_header_shading(table2.rows[0].cells[j], '1F4E79')

modules = [
    ('需求与设计', '需求规格、系统设计文档', '于枫硕', '彭仁希'),
    ('项目管理', '计划、进度、协调', '王启炀', '全体'),
    ('数据采集', '从网络/本地批量下载人脸素材', '侯贺升', '—'),
    ('人脸检测与对齐', '检测人脸位置、对齐裁剪为标准尺寸', '彭仁希', '吴迪'),
    ('质量评估', '模糊检测、光照检测、遮挡检测，筛选低质量图片', '吴迪', '彭仁希'),
    ('去重清洗', '感知哈希/特征匹配去重，清理无效素材', '侯贺升', '余野'),
    ('测试', '单元测试、集成测试、性能测试', '余野', '吴迪、侯贺升'),
    ('UI与展示', '清洗结果可视化、统计报表', '吴迪', '王启炀'),
]

for i, (mod, desc, owner, collab) in enumerate(modules):
    row = table2.rows[i + 1]
    cells = row.cells
    set_cell_text(cells[0], mod, font_size=Pt(10))
    set_cell_text(cells[1], desc, font_size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(cells[2], owner, font_size=Pt(10))
    set_cell_text(cells[3], collab, font_size=Pt(10))
    bg = 'EBF3FB' if i % 2 == 0 else 'F5F8FC'
    for c in cells:
        set_header_shading(c, bg)

# ════════════════════════════════════════════════
#  第四部分：开发排期
# ════════════════════════════════════════════════
doc.add_heading('四、开发排期', level=1)

phases = [
    ('第一阶段：需求与设计', '第1–2周', [
        '于枫硕 — 完成需求规格说明书（SRS）',
        '彭仁希 — 完成技术方案与架构设计',
        '王启炀 — 制定详细项目计划与里程碑',
    ]),
    ('第二阶段：核心开发', '第3–5周', [
        '侯贺升 — 完成数据采集模块与去重算法',
        '彭仁希 — 完成人脸检测与对齐模块',
        '吴迪 — 完成质量评估模块与前端界面',
        '于枫硕 — 完善需求文档，跟进开发进度',
    ]),
    ('第三阶段：集成与测试', '第6–7周', [
        '全体 — 集成各模块，联调接口',
        '余野 — 执行测试并输出测试报告',
        '吴迪 — 配合修复Bug，优化系统性能',
        '王启炀 — 跟踪进度，确保里程碑达成',
    ]),
    ('第四阶段：收尾与答辩', '第8周', [
        '全体 — 整理文档与代码，完善系统',
        '王启炀 — 主导答辩PPT与演示准备',
        '彭仁希 — 协助编写技术文档与部署说明',
        '全体 — 参与最终演示与答辩',
    ]),
]

for title, period, tasks in phases:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{title}（{period}）')
    run.font.bold = True
    run.font.size = Pt(11.5)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    for task in tasks:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(task)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── Footer note ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
run = p.add_run('— 本表为第五组内部任务分工，请各成员按计划按时完成 —')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── Save ──
output_path = r'C:\Users\q1948\Desktop\Course\软开3\第五组_任务分工分配表.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
